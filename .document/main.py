from random import random
import os
from docx import Document
from docx.shared import *

FOLDER_SEQUENCE = "./sequence"

def parse_mmdc_content(path: str) -> str:
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        title = lines[0].replace('#', '').strip()
        content = '\n'.join(lines[1:])
        return title, content

def render_mmd_to_img(mmd_content):
    os.makedirs('temp', exist_ok=True)

    random_prefix = str(random())
    random_input_file_name = os.path.join('temp', f'{random_prefix}.md')
    random_output_file_name = os.path.join('temp', f'{random_prefix}.png')
    
    with open(random_input_file_name, 'w', encoding='utf-8') as f:
        f.write(mmd_content)

    os.system(f'mmdc -i {random_input_file_name} -o {random_output_file_name}')

    return random_output_file_name

def process_sequence_diagram():
    document = Document('./template.docx')

    # Chỉnh layout thành a4
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    for topic in os.listdir(FOLDER_SEQUENCE):
        for mmd_file in os.listdir(os.path.join(FOLDER_SEQUENCE, topic)):
            path = os.path.join(FOLDER_SEQUENCE, topic, mmd_file)

            print(f'Rendering {path}')
            
            title, content = parse_mmdc_content(path)

            content = content.replace('```mermaid', '')
            content = content.replace('```', '')

            image_path = render_mmd_to_img(content)

            document.add_heading(title, 2)

            document.add_picture(image_path, Mm(120))

    os.makedirs('output', exist_ok=True)

    document.save(f'output/sequence-diagram.docx')

def clean():
    import shutil
    shutil.rmtree('./temp')

process_sequence_diagram()
clean()
